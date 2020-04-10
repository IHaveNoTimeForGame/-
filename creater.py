import requests,re,win32ui,logging,os
from tkinter import *
import tkinter.filedialog
from lxml import etree
import jieba
jieba.set_dictionary("./dict.txt")
jieba.initialize()
import jieba.analyse
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from random import choice
from time import time,sleep

def print_ssr():
    s = """
    *********************************************************************
    * 　　　　　　　 ┏┓       ┏┓+ +
    * 　　　　　　　┏┛┻━━━━━━━┛┻┓ + +
    * 　　　　　　　┃　　　　　　 ┃
    * 　　　　　　　┃　　　━　　　┃ ++ + + +
    * 　　　　　　 █████━█████  ┃+
    * 　　　　　　　┃　　　　　　 ┃ +
    * 　　　　　　　┃　　　┻　　　┃
    * 　　　　　　　┃　　　　　　 ┃ + +
    * 　　　　　　　┗━━┓　　　 ┏━┛
    *                      ┃　　  ┃
    * 　　　　　　　　　┃　　  ┃ + + + +
    * 　　　　　　　　　┃　　　┃　Code is far away from bug with the animal protecting
    * 　　　　　　　　　┃　　　┃ + 　　　　         神兽保佑, 代码无bug
    * 　　　　　　　　　┃　　　┃
    * 　　　　　　　　　┃　　　┃　　+
    * 　　　　　　　　　┃　 　 ┗━━━┓ + +
    * 　　　　　　　　　┃ 　　　　　┣┓
    * 　　　　　　　　　┃ 　　　　　┏┛
    * 　　　　　　　　　┗┓┓┏━━━┳┓┏┛ + + + +
    * 　　　　　　　　　 ┃┫┫　 ┃┫┫
    * 　　　　　　　　　 ┗┻┛　 ┗┻┛+ + + +
    ********************************************************************* """
    print(s)

def requests_ssr(url):
    is_succeed = False
    retry_count = 0
    max_retry_count = 10
    content = None
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/77.0.3865.90 Safari/537.36",
    }
    while is_succeed == False and retry_count < max_retry_count:
        try:
            response = requests.get(url,headers=headers)
            sleep(choice([0.2,0.4,0.6,0.8,1,1.5,2]))
            is_succeed = True
            content = response.content
        except Exception as e:
            retry_count += 1
            print(f"网络错误，正在重试第{retry_count}次...")
    return content

def get_text():
    root = Tk()
    root.withdraw()
    file_names = tkinter.filedialog.askopenfilenames(title='选择txt文本文件', filetypes=[('Text', '*.txt')])
    if len(file_names) != 0:
        string_filename = ""
        for i in file_names:
            string_filename += i + "\n"
        print("当前选择文件：\n",string_filename)
        for fn in file_names:
            with open(fn,encoding="utf-8") as f:
                s = f.read()
                create_docx(s)

def create_docx(s):
    document = Document()
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    os.makedirs("./pic",exist_ok=True)
    result = re.split("\n",s)
    for text in result:
        if text == "":
            pass
        else:
            print("--"*20)
            print("提取分段文本：",text)
            keywords = jieba.analyse.textrank(text, topK=3)
            keyword = ""
            if keywords == []:
                keyword = text
            else:
                keyword = choice(keywords)
            url = "https://www.52doutu.cn/search/" + keyword
            print("提取关键词：", keyword)
            content = requests_ssr(url)
            if content == None:
                print("请检查网络后，再重试...")
                break
            root = etree.HTML(content.decode())
            items = root.xpath('//div[@class="img-blocks"]/div[@class="img-block-item"]')
            if items == []:
                url = "https://www.52doutu.cn/search/" + "搞笑"
                content = requests_ssr(url)
                if content == None:
                    print("请检查网络后，再重试...")
                    break
                root = etree.HTML(content.decode())
                items = root.xpath('//div[@class="img-blocks"]/div[@class="img-block-item"]')
            item = choice(items)
            href = item.xpath("./a/@href")
            temp =str(time())
            with open(f"./pic/{temp}图片.jpg","wb") as f:
                content = requests_ssr(href[0])
                if content == None:
                    print("请检查网络后，再重试...")
                f.write(content)
            document.add_paragraph('').add_run(text + "\n")
            document.add_picture(f"./pic/{temp}图片.jpg",width=Inches(2))
        print("--" * 20)
    file_path = f'转换结果{time()}.docx'
    document.save(file_path)
    print(f"转换成功，输出文件名称：{file_path}")

while True:
    print_ssr()
    try:
        print("选择一个待转换的txt文件...")
        get_text()
    except Exception as e:
        print(e)
    input_txt = input("按q退出，其他键继续...")
    if input_txt == "q":
        break
    else:
        pass